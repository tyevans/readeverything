"""Word documents, OOXML and ODF.

A document's headings are its table of contents, and a table of contents is
what an agent needs in order to decide where to look — so the card's outline is
one `Segment` per heading and `represent` maps every character to the section
it fell under. Sections rather than pages: a `.docx` has no pages until
something lays it out, and claiming `PageRef` here would be inventing a
pagination the file does not carry.

Tables render as pipe-delimited rows IN DOCUMENT ORDER rather than being
skipped or appended, because a table is frequently the answer and where it sat
is part of what it means. That forces the walk over `body.iterchildren()` below
instead of the much simpler `document.paragraphs`, which silently omits every
paragraph inside a table and loses the position of every table.

`describe` parses the document, unlike `pdf.py`, which probes. There is no
cheaper probe that can answer "how many headings" — a heading count IS a parse
— and the parse costs no model call and no subprocess. A fabricated count would
be worse than a slightly more expensive card.

python-docx is imported directly here, guarded exactly as `pdf.py` guards
pypdfium2.
"""

from __future__ import annotations

import time
from dataclasses import dataclass
from typing import ClassVar

try:
    import docx
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph
except ImportError as exc:  # pragma: no cover - exercised via a patched sys.modules
    raise ImportError(
        "readeverything's Word support needs python-docx, which ships in the "
        "'office' extra: pip install 'readeverything[office]'. "
        "The composition root omits Word handling when python-docx is absent, so "
        "reaching this means the handler was imported directly."
    ) from exc
import io

from pydantic import BaseModel, Field

from readeverything.adapters.odf import odf_blocks
from readeverything.adapters.ooxml import ODF_TEXT_MIME, WORD_MIME, office_mimetype
from readeverything.domain.affordance import Affordance, DetailLevel
from readeverything.domain.capability import Capability
from readeverything.domain.card import Card, Segment
from readeverything.domain.errors import UnknownAffordanceError
from readeverything.domain.identity import MediaKind, SourceRef
from readeverything.domain.locator_map import LocatorMap, LocatorSegment
from readeverything.domain.locators import ByteRange, CharSpan
from readeverything.domain.observation import OperationFinished, OperationStarted
from readeverything.domain.rendition import (
    Budget,
    Degradation,
    Rendered,
    Rendition,
    TextContent,
)
from readeverything.ports.observation import Observer, emit
from readeverything.ports.source import SourceReader

#: What `represent` calls itself when it narrates, matching every other handler.
_OPERATION = "represent"

#: Every section's text ends with this, and the section's `LocatorSegment`
#: INCLUDES it. `LocatorMap` demands total, gapless, zero-start coverage and
#: `CharSpan.__post_init__` rejects `start >= end`, so a heading with an empty
#: body would otherwise contribute a zero-width span and raise. Owning the
#: separator means every section owns at least one character no matter what it
#: contained. Do not "simplify" this away: an empty section in the middle of a
#: document is what breaks the map.
SECTION_SEPARATOR = "\n\n"

#: What the section holding any text before the first heading is called. A
#: `LocatorMap` must start at offset 0, so there is no such thing as text
#: belonging to no section — but calling it a heading would put a title in the
#: table of contents that the author never wrote.
UNTITLED_SECTION = "(untitled)"

#: How a table's cells are joined. A row per line, cells separated by this.
CELL_DELIMITER = " | "

#: The style-name prefix python-docx gives every built-in heading level.
_HEADING_STYLE_PREFIX = "Heading"


@dataclass(frozen=True, slots=True)
class _Section:
    """One heading and the blocks beneath it, before any text is joined."""

    title: str
    is_titled: bool
    blocks: tuple[str, ...]

    def rendered(self) -> str:
        parts = ([self.title] if self.is_titled else []) + [b for b in self.blocks if b]
        return "\n".join(parts)


@dataclass(frozen=True, slots=True)
class _Counts:
    """The card's facts, gathered during the one parse that produces sections."""

    headings: int
    paragraphs: int
    words: int
    tables: int
    comments: int
    tracked: bool


class ReadSectionParams(BaseModel):
    index: int = Field(default=0, ge=0, description="0-indexed section number to read.")


class ReadRangeParams(BaseModel):
    start: int = Field(default=0, ge=0, description="First character to return.")
    end: int = Field(default=4096, gt=0, description="One past the last character to return.")


class ListCommentsParams(BaseModel):
    pass


class ReadTableParams(BaseModel):
    index: int = Field(default=0, ge=0, description="0-indexed table number to read.")


def _table_text(table: Table) -> str:
    return "\n".join(
        CELL_DELIMITER.join(" ".join(cell.text.split()) for cell in row.cells) for row in table.rows
    )


def _heading_level(paragraph: Paragraph) -> int:
    """The paragraph's heading level, or 0 when it is body text.

    Read from the style name rather than from the outline level, because that
    is what python-docx exposes and what `add_heading` writes. A style named
    `Heading 2` is level 2; `Title` and `Normal` are not headings.
    """
    name = paragraph.style.name if paragraph.style is not None else ""
    if not name.startswith(_HEADING_STYLE_PREFIX):
        return 0
    tail = name[len(_HEADING_STYLE_PREFIX) :].strip()
    try:
        return max(1, int(tail))
    except ValueError:
        return 1


class OfficeWordHandler:
    """Reads a Word document's body, and maps every character to its section."""

    mime_patterns: ClassVar[tuple[str, ...]] = (WORD_MIME, ODF_TEXT_MIME)
    priority: ClassVar[int] = 0
    handler_id: ClassVar[str] = "office_word"
    handler_version: ClassVar[int] = 1

    def __init__(self, *, source: SourceReader, observer: Observer | None = None) -> None:
        self._source = source
        self._observer = observer

    def requires(self) -> frozenset[Capability]:
        """Nothing. Reading a Word document needs no model and no binary."""
        return frozenset()

    def affordances(self) -> tuple[Affordance, ...]:
        return (
            Affordance(
                name="read_section",
                description="Return one section of the document: a heading and the text under it.",
                params=ReadSectionParams,
                requires=frozenset(),
                level=DetailLevel.SEGMENT,
            ),
            Affordance(
                name="read_range",
                description="Return a character range of the document's flattened text.",
                params=ReadRangeParams,
                requires=frozenset(),
                level=DetailLevel.SEGMENT,
            ),
            Affordance(
                name="list_comments",
                description="List every comment in the document, with its author.",
                params=ListCommentsParams,
                requires=frozenset(),
                level=DetailLevel.SEGMENT,
            ),
            Affordance(
                name="read_table",
                description="Return one table as pipe-delimited rows.",
                params=ReadTableParams,
                requires=frozenset(),
                level=DetailLevel.SEGMENT,
            ),
        )

    # -- parsing ---------------------------------------------------------

    def _is_odf(self, data: bytes) -> bool:
        return office_mimetype(data) == ODF_TEXT_MIME

    def _sections_from_odf(self, data: bytes) -> tuple[list[_Section], _Counts] | None:
        blocks = odf_blocks(data)
        if not blocks:
            return None
        sections: list[_Section] = []
        current: _Section | None = None
        pending: list[str] = []
        headings = 0
        paragraphs = 0
        for block in blocks:
            if block.level > 0:
                if current is not None:
                    sections.append(_Section(current.title, current.is_titled, tuple(pending)))
                elif pending:
                    sections.append(_Section(UNTITLED_SECTION, False, tuple(pending)))
                pending = []
                current = _Section(block.text, True, ())
                headings += 1
            else:
                paragraphs += 1
                pending.append(block.text)
        if current is not None:
            sections.append(_Section(current.title, current.is_titled, tuple(pending)))
        elif pending:
            sections.append(_Section(UNTITLED_SECTION, False, tuple(pending)))
        words = sum(len(b.text.split()) for b in blocks)
        return sections, _Counts(
            headings=headings,
            paragraphs=paragraphs,
            words=words,
            tables=0,
            comments=0,
            tracked=False,
        )

    def _sections_from_ooxml(self, data: bytes) -> tuple[list[_Section], _Counts] | None:
        try:
            document = docx.Document(io.BytesIO(data))
        except Exception:
            return None
        body = document.element.body
        sections: list[_Section] = []
        current_title: str | None = None
        pending: list[str] = []
        headings = 0
        paragraphs = 0
        tables = 0
        words = 0

        def close() -> None:
            nonlocal pending
            if current_title is not None:
                sections.append(_Section(current_title, True, tuple(pending)))
            elif pending:
                sections.append(_Section(UNTITLED_SECTION, False, tuple(pending)))
            pending = []

        try:
            for child in body.iterchildren():
                if child.tag == qn("w:p"):
                    paragraph = Paragraph(child, document)
                    text = " ".join(paragraph.text.split())
                    words += len(text.split())
                    if _heading_level(paragraph) > 0:
                        close()
                        current_title = text
                        headings += 1
                    else:
                        paragraphs += 1
                        pending.append(text)
                elif child.tag == qn("w:tbl"):
                    tables += 1
                    rendered = _table_text(Table(child, document))
                    words += len(rendered.split())
                    pending.append(rendered)
            close()
            comments = self._comment_lines(document)
            tracked = any(True for _ in body.iter(qn("w:ins"))) or any(
                True for _ in body.iter(qn("w:del"))
            )
        except Exception:
            return None
        return sections, _Counts(
            headings=headings,
            paragraphs=paragraphs,
            words=words,
            tables=tables,
            comments=len(comments),
            tracked=tracked,
        )

    def _comment_lines(self, document: object) -> tuple[str, ...]:
        """One line per comment, or an empty tuple when there are none.

        `getattr` rather than a direct attribute: `Document.comments` arrived in
        python-docx 1.2, and the `office` extra floors at 1.1. An older release
        must yield a document with no comments rather than an exception.
        """
        lines: list[str] = []
        for comment in getattr(document, "comments", ()):
            author = getattr(comment, "author", "") or "(unknown)"
            text = " ".join((getattr(comment, "text", "") or "").split())
            lines.append(f"{author}: {text}")
        return tuple(lines)

    def _parse(self, data: bytes) -> tuple[list[_Section], _Counts] | None:
        if self._is_odf(data):
            return self._sections_from_odf(data)
        return self._sections_from_ooxml(data)

    def _tables(self, data: bytes) -> tuple[str, ...]:
        if self._is_odf(data):
            return ()
        try:
            document = docx.Document(io.BytesIO(data))
            return tuple(_table_text(table) for table in document.tables)
        except Exception:
            return ()

    def _comments(self, data: bytes) -> tuple[str, ...]:
        if self._is_odf(data):
            return ()
        try:
            return self._comment_lines(docx.Document(io.BytesIO(data)))
        except Exception:
            return ()

    # -- flattening ------------------------------------------------------

    def _flatten(
        self, sections: list[_Section]
    ) -> tuple[str, tuple[LocatorSegment, ...], tuple[int, ...], tuple[Segment, ...]]:
        """The whole text, its map, its barriers, and the card's outline.

        Built once and shared, so the outline's `CharSpan` addresses exactly the
        text `represent` returns. Computing them separately is how a table of
        contents acquires the wrong page numbers.
        """
        chunks: list[str] = []
        segments: list[LocatorSegment] = []
        barriers: list[int] = []
        outline: list[Segment] = []
        cursor = 0
        for index, section in enumerate(sections):
            chunk = section.rendered() + SECTION_SEPARATOR
            if index:
                barriers.append(cursor)
            span = CharSpan(cursor, cursor + len(chunk))
            segments.append(LocatorSegment(span, span))
            if section.is_titled:
                outline.append(Segment(span, section.title))
            cursor += len(chunk)
            chunks.append(chunk)
        return "".join(chunks), tuple(segments), tuple(barriers), tuple(outline)

    # -- the handler surface ---------------------------------------------

    async def describe(self, ref: SourceRef) -> Card:
        """Shape and table of contents, from a parse rather than from a probe.

        `kind` is `BINARY`, not some new `DOCUMENT` member: these mimetypes
        reach this handler at the registry's exact-mimetype step, long before
        the kind step, so what identifies a Word card is its mime, its counts
        and its affordances.
        """
        data = await self._source.read_bytes(ref.uri)
        parsed = self._parse(data)
        if parsed is None:
            return Card(
                ref=ref,
                kind=MediaKind.BINARY,
                facts={"readable": "no", "size_bytes": ref.size_bytes},
                outline=(),
                excerpt=None,
                affordances=self.affordances(),
            )
        sections, counts = parsed
        _text, _segments, _barriers, outline = self._flatten(sections)
        return Card(
            ref=ref,
            kind=MediaKind.BINARY,
            facts={
                "readable": "yes",
                "heading_count": counts.headings,
                "paragraph_count": counts.paragraphs,
                "word_count": counts.words,
                "table_count": counts.tables,
                "comment_count": counts.comments,
                "tracked_changes": "yes" if counts.tracked else "no",
                "size_bytes": ref.size_bytes,
            },
            outline=outline,
            excerpt=None,
            affordances=self.affordances(),
        )

    async def invoke(self, ref: SourceRef, name: str, params: BaseModel) -> Rendition:
        match name:
            case "read_section":
                if not isinstance(params, ReadSectionParams):
                    raise TypeError(f"expected ReadSectionParams, got {type(params).__name__}")
                return await self._read_section(ref, params.index)
            case "read_range":
                if not isinstance(params, ReadRangeParams):
                    raise TypeError(f"expected ReadRangeParams, got {type(params).__name__}")
                return await self._read_range(ref, params)
            case "list_comments":
                if not isinstance(params, ListCommentsParams):
                    raise TypeError(f"expected ListCommentsParams, got {type(params).__name__}")
                return await self._list_comments(ref)
            case "read_table":
                if not isinstance(params, ReadTableParams):
                    raise TypeError(f"expected ReadTableParams, got {type(params).__name__}")
                return await self._read_table(ref, params.index)
            case _:
                raise UnknownAffordanceError(name, (a.name for a in self.affordances()))

    def _degraded(self, ref: SourceRef, detail: str) -> Rendition:
        """What every unreadable or out-of-range request returns.

        Never an exception: an agent guessing a section index gets a result it
        can read and correct.
        """
        return Rendition(
            locator=ByteRange(0, max(1, ref.size_bytes)),
            content=TextContent(detail),
            degraded=True,
        )

    async def _read_section(self, ref: SourceRef, index: int) -> Rendition:
        data = await self._source.read_bytes(ref.uri)
        parsed = self._parse(data)
        if parsed is None:
            return self._degraded(ref, f"{ref.uri} could not be opened as a Word document")
        sections, _counts = parsed
        if index >= len(sections):
            return self._degraded(
                ref, f"section {index} does not exist; the document has {len(sections)} section(s)"
            )
        _text, _segments, _barriers, _outline = self._flatten(sections)
        chunk = sections[index].rendered() + SECTION_SEPARATOR
        start = sum(len(s.rendered() + SECTION_SEPARATOR) for s in sections[:index])
        return Rendition(
            locator=CharSpan(start, start + len(chunk)),
            content=TextContent(sections[index].rendered()),
        )

    async def _read_range(self, ref: SourceRef, params: ReadRangeParams) -> Rendition:
        data = await self._source.read_bytes(ref.uri)
        parsed = self._parse(data)
        if parsed is None:
            return self._degraded(ref, f"{ref.uri} could not be opened as a Word document")
        sections, _counts = parsed
        text, _segments, _barriers, _outline = self._flatten(sections)
        # Clamped rather than rejected: an agent that asked for more than the
        # document holds gets what there is, which is the answer it wanted.
        end = min(params.end, len(text))
        start = min(params.start, max(0, end - 1))
        if end <= start:
            return self._degraded(ref, "the document has no text in that range")
        return Rendition(locator=CharSpan(start, end), content=TextContent(text[start:end]))

    async def _list_comments(self, ref: SourceRef) -> Rendition:
        data = await self._source.read_bytes(ref.uri)
        comments = self._comments(data)
        if not comments:
            # "There are none" and empty output are different answers, and only
            # one of them tells an agent to stop looking.
            return Rendition(
                locator=ByteRange(0, max(1, ref.size_bytes)),
                content=TextContent("the document carries no comments"),
            )
        return Rendition(
            locator=ByteRange(0, max(1, ref.size_bytes)),
            content=TextContent("\n".join(comments)),
        )

    async def _read_table(self, ref: SourceRef, index: int) -> Rendition:
        data = await self._source.read_bytes(ref.uri)
        tables = self._tables(data)
        if index >= len(tables):
            return self._degraded(
                ref, f"table {index} does not exist; the document has {len(tables)} table(s)"
            )
        return Rendition(
            locator=ByteRange(0, max(1, ref.size_bytes)),
            content=TextContent(tables[index]),
        )

    async def represent(self, ref: SourceRef, budget: Budget) -> Rendered:
        """Narrated start to finish, matching every other handler."""
        emit(self._observer, OperationStarted(operation=_OPERATION, ref=ref))
        started = time.perf_counter()
        try:
            return await self._represent(ref, budget)
        finally:
            emit(
                self._observer,
                OperationFinished(
                    operation=_OPERATION, ref=ref, elapsed_s=time.perf_counter() - started
                ),
            )

    async def _represent(self, ref: SourceRef, budget: Budget) -> Rendered:
        data = await self._source.read_bytes(ref.uri)
        parsed = self._parse(data)
        if parsed is None:
            return self._nothing_to_read(
                ref,
                budget,
                summary=f"Unreadable Word document {ref.uri}, {ref.size_bytes} bytes.",
                what="document unopenable",
                detail="the file could not be opened as a Word document; no text was extracted",
            )
        sections, _counts = parsed
        if not sections:
            return self._nothing_to_read(
                ref,
                budget,
                summary=f"Word document {ref.uri} has no body text.",
                what="document has no body",
                detail="the document opened but contains no paragraphs; no text was extracted",
            )
        text, segments, barriers, _outline = self._flatten(sections)
        return self._fit(text, segments, barriers, budget, ())

    def _nothing_to_read(
        self, ref: SourceRef, budget: Budget, *, summary: str, what: str, detail: str
    ) -> Rendered:
        """A rendition for a file with no section to point at.

        Located by `ByteRange` rather than `CharSpan` into a body: no section
        was ever observed, and citing one would be a claim about a document
        this handler never read.
        """
        segments = (
            LocatorSegment(CharSpan(0, len(summary)), ByteRange(0, max(1, ref.size_bytes))),
        )
        return self._fit(summary, segments, (), budget, (Degradation(what=what, detail=detail),))

    def _fit(
        self,
        full: str,
        segments: tuple[LocatorSegment, ...],
        barriers: tuple[int, ...],
        budget: Budget,
        degradations: tuple[Degradation, ...],
    ) -> Rendered:
        """Apply the budget, pruning the map and the barriers along with the text.

        `Rendered` rejects a map that does not cover its text exactly and a
        barrier past the end, so truncation cannot touch the text alone. A
        budget of zero still keeps one character, because `CharSpan(0, 0)`
        raises, and the degradation reports the character kept rather than the
        budget asked for.
        """
        if budget.max_chars is None or len(full) <= budget.max_chars:
            return Rendered(
                text=full,
                locator_map=LocatorMap.build(segments),
                barriers=barriers,
                degradations=degradations,
            )
        keep = max(1, budget.max_chars)
        text = full[:keep]
        kept = tuple(
            LocatorSegment(CharSpan(s.span.start, min(s.span.end, keep)), s.locator)
            for s in segments
            if s.span.start < keep
        )
        return Rendered(
            text=text,
            locator_map=LocatorMap.build(kept),
            barriers=tuple(barrier for barrier in barriers if barrier < keep),
            degradations=(
                *degradations,
                Degradation(
                    what="text truncated",
                    detail=f"kept {len(text)} of {len(full)} characters",
                ),
            ),
        )
